# views.py
from django.contrib.auth import authenticate, login, logout
from .forms import UserCreationForm, LoginForm
from difflib import SequenceMatcher
from django.shortcuts import render, redirect, HttpResponse
from django.contrib.auth.decorators import login_required
from .models import FileKeywordCount
import os
import math
import PyPDF2
from docx import Document
from openpyxl import load_workbook
from .forms import FileUploadForm
from .utils import extract_pdf_preview, extract_docx_preview, extract_xlsx_preview
from .models import AcceptScore
from .models import KeyWord
from .forms import DocumentForm
from .models import AdminInput, UploadedFile
from textblob import TextBlob
import docx
import openpyxl
import PyPDF2
import nltk
nltk.download('vader_lexicon')
from PyPDF2 import PdfReader
import io
from docx import Document
import re
import nltk
from nltk.sentiment import SentimentIntensityAnalyzer
from django.utils.safestring import mark_safe
from django.utils import timezone
from django.urls import reverse
from django.conf import settings
# Initialize the sentiment analyzer
nltk.download('vader_lexicon')
sid = SentimentIntensityAnalyzer()

from django.contrib.auth.decorators import login_required
from django.contrib import messages
from .forms import UserEditForm


def test(request):
    keywords = KeyWord.objects.all()

    # Retrieve AcceptScore data from the database
    accept_scores = AcceptScore.objects.all()

    return render(request, 'test.html', {'keywords': keywords, 'accept_scores': accept_scores})

@login_required
def Profile(request):
    if request.method == 'POST':
        form = UserEditForm(request.POST, instance=request.user)
        if form.is_valid():
            user = form.save(commit=False)
            old_password = form.cleaned_data.get('old_password')
            new_password = form.cleaned_data.get('new_password')

            if old_password and new_password:
                user.set_password(new_password)
                messages.success(request, 'Password updated successfully.')
            user.save()
            messages.success(request, 'Profile updated successfully.')
            return redirect('Profile')  # Change 'profile' to the actual URL pattern name for the profile view
    else:
        form = UserEditForm(instance=request.user)
    return render(request, 'Profile.html', {'form': form})

# Home page
def index(request):
    # Get all keywords from the database
    keywords = KeyWord.objects.all()

    # Retrieve AcceptScore data from the database
    accept_scores = AcceptScore.objects.all()

    return render(request, 'index.html', {'keywords': keywords, 'accept_scores': accept_scores})



# signup page
def user_signup(request):
    if request.method == 'POST':
        form = UserCreationForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('login')
    else:
        form = UserCreationForm()
    return render(request, 'signup.html', {'form': form})



# login page
def user_login(request):
    if request.method == 'POST':
        form = LoginForm(request.POST)
        if form.is_valid():
            username = form.cleaned_data['username']
            password = form.cleaned_data['password']
            user = authenticate(request, username=username, password=password)
            if user:
                login(request, user)
                return redirect('test')
    else:
        form = LoginForm()
    return render(request, 'login.html', {'form': form})



# logout page
def user_logout(request):
    logout(request)
    return redirect('login')


@login_required
def keyword_scan(request):
    if request.method == 'POST' and request.FILES['file']:
        uploaded_file = request.FILES['file']
        allowed_extensions = ['.pdf', '.docx', '.xlsx']
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()

        if file_extension not in allowed_extensions:
            return render(request, 'file_upload_error.html',
                          {'error': 'Unsupported file format. Please upload a PDF, DOCX, or XLSX file.'})

        form = FileUploadForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = request.FILES['file']
            file_extension = uploaded_file.name.split('.')[-1].lower()
            preview_content = None

            # Read the uploaded file and extract text
            if file_extension == 'pdf':
                text = read_pdf(uploaded_file)
                preview_content = extract_pdf_preview(uploaded_file)
            elif file_extension == 'docx':
                text = read_docx(uploaded_file)
                preview_content = extract_docx_preview(uploaded_file)
            elif file_extension == 'xlsx':
                text = read_xlsx(uploaded_file)
                preview_content = extract_xlsx_preview(uploaded_file)
            else:
                text = ''

            # Pass file name to template context
            file_name = uploaded_file.name

            # Get all keywords from the database
            keywords = KeyWord.objects.all()

            # Retrieve AcceptScore data from the database
            accept_scores = AcceptScore.objects.all()

            # Count occurrences of keywords in the file content and get their scores
            keyword_count = {}
            overall_total = 0  # Initialize overall total
            for keyword in keywords:
                count = 0  # Initialize count for each keyword
                for paragraph in text.split('\n\n'):  # Split text into paragraphs
                    count += len(re.findall(r'\b{}\b'.format(re.escape(keyword.keywords)), paragraph, re.IGNORECASE))
                if count > 0:
                    # Calculate total by multiplying count with score
                    total = count * keyword.score
                    keyword_count[keyword.keywords] = {'count': count, 'score': keyword.score, 'total': total}
                    overall_total += total  # Add total to overall total

            # Check if overall_total meets the accept score criteria
            accept_score_value = accept_scores.first().score  # Assuming there's only one accept score
            score_description = ""
            if overall_total == accept_score_value:
                score_description = "Score Passed"
            elif overall_total < accept_score_value:
                score_description = "Failed, Score is Below the passing score"
            else:
                score_description = "Failed, Score is Higher than the passing score"

            # Store file and keyword count in the database
            file_entry = FileKeywordCount.objects.create(
                user=request.user,
                file=uploaded_file,
                file_type=file_extension,
                keyword_count=keyword_count
            )

            return render(request, 'index.html', {'keyword_count': keyword_count, 'preview_content': preview_content,
                                                  'overall_total': overall_total, 'file_name': file_name,
                                                  'accept_scores': accept_scores,
                                                  'score_description': score_description, 'keywords': keywords})

    else:
        form = FileUploadForm()
    return render(request, 'keyword_scan.html', {'form': form})


def read_pdf(uploaded_file):
    reader = PyPDF2.PdfReader(uploaded_file)
    text = ''
    for page in range(len(reader.pages)):
        text += reader.pages[page].extract_text()
    return text


def read_docx(uploaded_file):
    doc = Document(uploaded_file)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    return text

def read_docx(uploaded_file):
    docx_content = io.BytesIO(uploaded_file.read())
    doc = Document(docx_content)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    return text


def read_xlsx(uploaded_file):
    wb = load_workbook(uploaded_file)
    text = ''
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        for row in ws.iter_rows():
            for cell in row:
                text += str(cell.value) + ' '
    return text


@login_required
def user_files(request):
    # Retrieve all files uploaded by the current user
    user_files = FileKeywordCount.objects.filter(user=request.user)

    Sentiment_file = UploadedFile.objects.filter(user=request.user)

    # Check if user_files has any data
    if user_files.exists():
        # Calculate overall score and remarks for each file
        for file_entry in user_files:
            overall_total = sum(data['total'] for data in file_entry.keyword_count.values())
            accept_score = AcceptScore.objects.first()  # Assuming only one accept score exists

            if overall_total == accept_score.score:
                score_description = "Score Passed"
            elif overall_total < accept_score.score:
                score_description = "Failed, Score is Below the passing score"
            else:
                score_description = "Failed, Score is Higher than the passing score"

            file_entry.overall_total = overall_total
            file_entry.score_description = score_description
    else:
        # If user_files is empty, return a string indicating no data available
        return render(request, 'user_files.html', {'no_data_available': True, 'Sentiment_file': Sentiment_file})

    # Check if Sentiment_file has any data
    if Sentiment_file.exists():
        # If Sentiment_file contains data, proceed with rendering the template
        return render(request, 'user_files.html', {'user_files': user_files, 'Sentiment_file': Sentiment_file})
    else:
        # If Sentiment_file is empty, return a string indicating no data available for sentiment analysis files
        return render(request, 'user_files.html', {'no_sentiment_data_available': True, 'user_files': user_files})

#Function For BackUp feautures
@login_required
def backUp(request):
    # Retrieve all files uploaded by the current user
    user_files = FileKeywordCount.objects.filter(user=request.user)
    Sentiment_file = UploadedFile.objects.filter(user=request.user)
    return render(request, 'backUp.html', {'user_files': user_files, 'Sentiment_file': Sentiment_file})

@login_required
def move_file(request, file_id):
    if request.method == 'POST':
        try:
            # Check if the object is an instance of FileKeywordCount
            if UploadedFile.objects.filter(pk=file_id).exists():
                # If so, update fields for UploadedFile
                file_instance = UploadedFile.objects.get(pk=file_id)
                file_instance.backUp = 2
                file_instance.date_backUp = timezone.now()
                # Update fields for UploadedFile if needed
                file_instance.save()
            else:
                # Handle the case where the object doesn't belong to either model
                return HttpResponse("File not found or unknown file type.")
        except ObjectDoesNotExist:
            # Handle the case where the object does not exist
            return HttpResponse("File does not exist.")
        # Redirect to a page displaying user files
        return redirect('user_files')
    else:
        # Handle other HTTP methods, e.g., GET
        return HttpResponseNotAllowed(['POST'])
    
@login_required
def move_fileKey(request, file_id):
    if request.method == 'POST':
        try:
            # Check if the object is an instance of FileKeywordCount
            if FileKeywordCount.objects.filter(pk=file_id).exists():
                # If so, update fields for FileKeywordCount
                file_instance = FileKeywordCount.objects.get(pk=file_id)
                file_instance.backUp = 2
                file_instance.date_backUp = timezone.now()
                file_instance.save()
            else:
                # Handle the case where the object doesn't belong to either model
                return HttpResponse("File not found or unknown file type.")
        except ObjectDoesNotExist:
            # Handle the case where the object does not exist
            return HttpResponse("File does not exist.")
        # Redirect to a page displaying user files
        return redirect('user_files')
    else:
        # Handle other HTTP methods, e.g., GET
        return HttpResponseNotAllowed(['POST'])


@login_required
def restore_file(request, file_id):
    try:

            # If so, update fields for FileKeywordCount
            file_entry = FileKeywordCount.objects.get(pk=file_id)
            file_entry.backUp = 1
            file_entry.date_backUp = timezone.now()
            file_entry.save()

    except  :
        # Handle the case where the object does not exist
        return HttpResponse("File does not exist.")

    # Redirect the user to a specific URL after restoring the file
    return redirect(reverse('backUp'))

def restore_files(request, file_id):
    try:

        # If so, update fields for UploadedFile
        file_entry = UploadedFile.objects.get(pk=file_id)
        # Update fields for UploadedFile
        # For example, if you want to set a default value for file_type
        file_entry.backUp = 1
        file_entry.date_backUp = timezone.now()
        file_entry.file_type = "default_file_type"
        file_entry.save()
       
    except  :
        # Handle the case where the object does not exist
        return HttpResponse("File does not exist.")

    # Redirect the user to a specific URL after restoring the file
    return redirect(reverse('backUp'))



# file deletion
@login_required
def delete_file(request, file_id):
    try:

        file_entry = FileKeywordCount.objects.get(id=file_id)
        is_keyword_count = True


        # Check if the logged-in user owns the file
        if file_entry.user == request.user:

            # For FileKeywordCount
            file_path = os.path.join(settings.MEDIA_ROOT, file_entry.file.name)

                
            if os.path.exists(file_path):
                os.remove(file_path)

            # Delete the file entry from the database
            file_entry.delete()

            # Redirect back to the backup page
            return redirect('backUp')
        else:
            # Return an error message or handle unauthorized access
            return HttpResponse("You are not authorized to delete this file.")
    except (FileKeywordCount.DoesNotExist):
        # Handle the case where the file entry does not exist
        return HttpResponse("File entry does not exist.")
    except Exception as e:
        # Handle any other exceptions
        return HttpResponse(str(e))
    
@login_required
def delete_files(request, file_id):
    try:

        file_entry = UploadedFile.objects.get(id=file_id)
        is_keyword_count = False

        # Check if the logged-in user owns the file
        if file_entry.user == request.user:
            # Delete the file from the local storage

            # For UploadedFile
            file_path = os.path.join(settings.MEDIA_ROOT, file_entry.document.name)
                
            if os.path.exists(file_path):
                os.remove(file_path)

            # Delete the file entry from the database
            file_entry.delete()

            # Redirect back to the backup page
            return redirect('backUp')
        else:
            # Return an error message or handle unauthorized access
            return HttpResponse("You are not authorized to delete this file.")
    except (UploadedFile.DoesNotExist):
        # Handle the case where the file entry does not exist
        return HttpResponse("File entry does not exist.")
    except Exception as e:
        # Handle any other exceptions
        return HttpResponse(str(e))



@login_required
def file_details(request, file_id):
    try:
        # Retrieve the file details based on the file_id
        file_details = FileKeywordCount.objects.get(id=file_id)

        # Retrieve the accept score
        accept_score = AcceptScore.objects.first()  # Assuming only one accept score exists

        # Determine the remarks based on overall score and accept score
        overall_total = sum(data['total'] for data in file_details.keyword_count.values())

        score_description = ""
        if overall_total == accept_score.score:
            score_description = "Score Passed"
        elif overall_total < accept_score.score:
            score_description = "Failed, Score is Below the passing score"
        else:
            score_description = "Failed, Score is Higher than the passing score"

        return render(request, 'file_details.html', {'file_details': file_details,
                                                     'overall_total': overall_total,
                                                     'accept_scores': accept_score,
                                                     'score_description': score_description})

    except FileKeywordCount.DoesNotExist:
        # Handle the case where the file does not exist
        return HttpResponse("File does not exist.")



def back_or_default(request, default_url='/'):
    """
    Redirects the user to the previous page they were on before accessing
    the current view. If the HTTP_REFERER header is not present or empty,
    it redirects to the specified default URL.
    """
    previous_page = request.META.get('HTTP_REFERER')
    if previous_page:
        return redirect(previous_page)
    else:
        return redirect(default_url)






##### SENTIMENT FUNCTION #####

def analyze_sentiment(text1, text2):
    # Analyze sentiment using TextBlob
    sentiment1 = TextBlob(text1).sentiment.polarity
    sentiment2 = TextBlob(text2).sentiment.polarity

    # Calculate the absolute difference between the sentiment polarity scores
    diff = abs(sentiment1 - sentiment2)

    # Map the difference to a score between 0 and 100
    similarity_score = 100 - (diff * 100)

    # Ensure the score is within the valid range [0, 100]
    similarity_score = max(0, min(100, similarity_score))

    # Return the similarity score
    return similarity_score

def extract_text_from_pdf(file):
    text = ""
    try:
        # Open the PDF file
        with open(file, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfFileReader(pdf_file)

            # Iterate through each page of the PDF
            for page_num in range(pdf_reader.numPages):
                page = pdf_reader.getPage(page_num)
                text += page.extractText()

        return text
    except Exception as e:
        # Handle any errors that may occur during text extraction
        print(f"Error extracting text from PDF: {e}")
        return ""

def extract_text_from_docx(file):
    doc = docx.Document(file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text
    return text

def extract_text_from_xlsx(file):
    wb = openpyxl.load_workbook(file)
    text = ""
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        for row in ws.iter_rows(values_only=True):
            text += " ".join(str(cell) for cell in row if cell is not None) + "\n"
    return text


def admin_input(request):
    if request.method == 'POST':
        paragraph = request.POST.get('paragraph')
        admin_input = AdminInput.objects.create(paragraph=paragraph)
        return redirect('user_upload')
    return render(request, 'admin_input.html')


def create_preview(file): # PREVIEW FORMATING
    if file.name.endswith('.docx'):
        doc = Document(io.BytesIO(file.read()))
        preview_text = ""
        for paragraph in doc.paragraphs:
            sentences = nltk.sent_tokenize(paragraph.text.strip())  # Tokenize paragraph into sentences
            for sentence in sentences:
                preview_text += sentence.strip() + '\n'  # Add each sentence followed by a newline
            preview_text += "    " + sentence.strip() + "    \n\n"  # Add a newline after each paragraph
        return preview_text
    elif file.name.endswith('.pdf'):
        reader = PdfReader(io.BytesIO(file.read()))
        preview_text = ""
        for page in reader.pages:
            text = page.extract_text()
            lines = text.split('\n')
            for line in lines:
                preview_text += "    " + line.strip() + "    \n\n"  # Add indentation on both left and right sides
        return preview_text

    elif file.name.endswith('.xlsx'):
        wb = openpyxl.load_workbook(io.BytesIO(file.read()))
        preview_text = ""
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            for row in sheet.iter_rows():
                for cell in row:
                    preview_text += "\t" + str(cell.value) + " "
                preview_text += "\n"
        return preview_text
    else:
        # Handle other file types accordingly
        return ""


def compare_paragraphs(admin_paragraphs, paragraph):
    max_similarity_score = 0
    
    # Iterate over each admin paragraph
    for admin_paragraph in admin_paragraphs:
        # Perform sentiment analysis on both paragraphs
        similarity_score = analyze_sentiment(admin_paragraph, paragraph)
        
        # Update the maximum similarity score if needed
        max_similarity_score = max(max_similarity_score, similarity_score)
    
    return max_similarity_score


def create_preview_with_highlights(file, admin_paragraph):
    if file.name.endswith('.docx'):
        doc = Document(io.BytesIO(file.read()))
        preview_text = ""
        for paragraph in doc.paragraphs:
            # Check similarity with admin paragraph
            similarity_score = compare_paragraphs(admin_paragraph, paragraph.text.strip())
            if similarity_score > 85:  # Example threshold, adjust as needed
                preview_text += '<span class="highlight">{}</span>'.format(paragraph.text.strip()) + '\n'
            else:
                preview_text += paragraph.text.strip() + '\n'
        return preview_text
    elif file.name.endswith('.pdf'):
        reader = PdfReader(io.BytesIO(file.read()))
        preview_text = ""
        for page in reader.pages:
            text = page.extract_text()
            lines = text.split('\n')
            for line in lines:
                # Check similarity with admin paragraph
                similarity_score = compare_paragraphs(admin_paragraph, line.strip())
                if similarity_score > 99:  # Example threshold, adjust as needed
                    preview_text += "<span class='highlight'>{}</span>".format(line.strip()) + '\n'
                else:
                    preview_text += line.strip() + '\n'
        return preview_text
    elif file.name.endswith('.xlsx'):
        wb = openpyxl.load_workbook(io.BytesIO(file.read()))
        preview_text = ""
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            for row in sheet.iter_rows():
                for cell in row:
                    # Check similarity with admin paragraph
                    similarity_score = compare_paragraphs(admin_paragraph, str(cell.value))
                    if similarity_score > 99:  # Example threshold, adjust as needed
                        preview_text += "<span class='highlight'>{}</span>".format(
                            str(cell.value)) + " "
                    else:
                        preview_text += str(cell.value) + " "
                preview_text += "\n"
        return preview_text
    else:
        # Handle other file types accordingly
        return ""


from itertools import zip_longest


def user_upload(request):
    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = request.FILES['document']
            
            # Extract the file extension
            file_extension = uploaded_file.name.split('.')[-1]

            # Get the latest admin paragraph
            admin_paragraph = AdminInput.objects.last().paragraph

            # Create preview content with highlights
            uploaded_text_preview = create_preview_with_highlights(uploaded_file, admin_paragraph)

            # Get all admin paragraphs
            admin_paragraphs = AdminInput.objects.all().values_list('paragraph', flat=True)

            # Initialize lists to store similarity scores and admin paragraphs
            admin_similarity_scores = []
            admin_paragraph_list = []

            # Compare the similarity of paragraphs based on sentiment analysis
            for admin_paragraph, uploaded_paragraph in zip(admin_paragraphs, uploaded_text_preview.split('\n')):
                similarity_score = compare_paragraphs(admin_paragraph, uploaded_paragraph)
                admin_similarity_scores.append(round(similarity_score, 2))
                admin_paragraph_list.append(admin_paragraph)

            # Round the overall similarity score to 2 decimal points
            average_similarity_score = sum(admin_similarity_scores) / len(admin_similarity_scores)
            average_similarity_score = round(average_similarity_score, 2)

            # Save the uploaded file and similarity score to the database
            uploaded_file_obj = UploadedFile(user=request.user, document=uploaded_file,
                                             similarity_score=average_similarity_score,
                                             file_type=file_extension)
            uploaded_file_obj.save()

            # Mark the preview content as safe for HTML rendering
            uploaded_text_preview = mark_safe(uploaded_text_preview)

            # Zip admin paragraphs and similarity scores
            admin_data = zip(admin_paragraph_list, admin_similarity_scores)

            return render(request, 'result.html', {'average_similarity_score': average_similarity_score,
                                                    'uploaded_file_obj': uploaded_file_obj,
                                                    'uploaded_text_preview': uploaded_text_preview,
                                                    'admin_data': admin_data,
                                                    'file_extension': file_extension})

    else:
        form = DocumentForm()
    return render(request, 'sentiment_forms.html', {'form': form})
